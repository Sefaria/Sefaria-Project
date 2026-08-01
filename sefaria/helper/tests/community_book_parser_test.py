import pytest
from io import BytesIO
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from sefaria.helper.community_book_parser import (
    ParseError, ParseResult, Chapter, Section,
    parse_document, build_jagged_array, build_schema,
)


def _make_docx(paragraphs):
    """Build an in-memory .docx from a list of (style, text) pairs."""
    doc = DocxDocument()
    for style, text in paragraphs:
        doc.add_paragraph(text, style=style)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class TestParseDocumentEntryPoint:
    def test_unsupported_file_type_raises(self):
        with pytest.raises(ParseError, match="Unsupported file type"):
            parse_document(None, "txt", 1)

    def test_pdf_raises_not_yet_supported(self):
        with pytest.raises(ParseError, match="PDF upload is not yet supported"):
            parse_document(None, "pdf", 1)


class TestBuildSchema:
    def test_depth_1_schema(self):
        schema = build_schema(1, "Test Book", "ספר בדיקה")
        assert schema["nodeType"] == "JaggedArrayNode"
        assert schema["depth"] == 1
        assert schema["addressTypes"] == ["Integer"]
        assert schema["sectionNames"] == ["Chapter"]
        titles = schema["titles"]
        en_primary = [t for t in titles if t.get("primary") and t["lang"] == "en"]
        assert len(en_primary) == 1
        assert en_primary[0]["text"] == "Test Book"

    def test_depth_2_schema(self):
        schema = build_schema(2, "Test Book", "ספר בדיקה")
        assert schema["depth"] == 2
        assert schema["addressTypes"] == ["Integer", "Integer"]
        assert schema["sectionNames"] == ["Chapter", "Section"]


def _array_nesting_depth(arr):
    if not isinstance(arr, list):
        return 0
    if not arr:
        return 1
    return 1 + _array_nesting_depth(arr[0])


class TestBuildJaggedArray:
    def test_depth_1_array(self):
        result = ParseResult(
            chapters=[
                Chapter(title="Ch1", sections=[], content=["para1", "para2"], word_count=2),
                Chapter(title="Ch2", sections=[], content=["para3"], word_count=1),
            ],
            total_words=3,
            structure_depth=1,
        )
        ja = build_jagged_array(result)
        assert ja == ["para1<br>para2", "para3"]

    def test_depth_2_array(self):
        result = ParseResult(
            chapters=[
                Chapter(
                    title="Ch1",
                    sections=[
                        Section(title="S1", content=["a", "b"], word_count=2),
                        Section(title="S2", content=["c"], word_count=1),
                    ],
                    content=[],
                    word_count=3,
                ),
            ],
            total_words=3,
            structure_depth=2,
        )
        ja = build_jagged_array(result)
        assert ja == [["a<br>b", "c"]]

    def test_depth_1_nesting_matches_structure_depth(self):
        result = ParseResult(
            chapters=[
                Chapter(title="Ch1", sections=[], content=["para1", "para2"], word_count=2),
                Chapter(title="Ch2", sections=[], content=["para3"], word_count=1),
            ],
            total_words=3,
            structure_depth=1,
        )
        ja = build_jagged_array(result)
        assert _array_nesting_depth(ja) == result.structure_depth

    def test_depth_2_nesting_matches_structure_depth(self):
        result = ParseResult(
            chapters=[
                Chapter(
                    title="Ch1",
                    sections=[
                        Section(title="S1", content=["a", "b"], word_count=2),
                        Section(title="S2", content=["c"], word_count=1),
                    ],
                    content=[],
                    word_count=3,
                ),
            ],
            total_words=3,
            structure_depth=2,
        )
        ja = build_jagged_array(result)
        assert _array_nesting_depth(ja) == result.structure_depth


class TestParseDocxDepth1:
    def test_valid_depth1_two_chapters(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Normal", "First paragraph."),
            ("Normal", "Second paragraph."),
            ("Heading 2", "Chapter Two"),
            ("Normal", "Third paragraph."),
        ])
        result = parse_document(buf, "docx", 1)
        assert len(result.chapters) == 2
        assert result.chapters[0].title == "Chapter One"
        assert result.chapters[0].content == ["First paragraph.", "Second paragraph."]
        assert result.chapters[0].word_count == 4
        assert result.chapters[1].title == "Chapter Two"
        assert result.chapters[1].content == ["Third paragraph."]
        assert result.structure_depth == 1

    def test_markdown_style_headings_fallback(self):
        buf = _make_docx([
            ("Normal", "## Chapter A"),
            ("Normal", "Some content here."),
            ("Normal", "## Chapter B"),
            ("Normal", "More content."),
        ])
        result = parse_document(buf, "docx", 1)
        assert len(result.chapters) == 2
        assert result.chapters[0].title == "Chapter A"
        assert result.chapters[1].title == "Chapter B"


class TestParseDocxDepth2:
    def test_valid_depth2(self):
        buf = _make_docx([
            ("Heading 2", "Part One"),
            ("Heading 3", "Section 1.1"),
            ("Normal", "Content for 1.1."),
            ("Heading 3", "Section 1.2"),
            ("Normal", "Content for 1.2."),
            ("Heading 2", "Part Two"),
            ("Heading 3", "Section 2.1"),
            ("Normal", "Content for 2.1."),
        ])
        result = parse_document(buf, "docx", 2)
        assert len(result.chapters) == 2
        assert result.chapters[0].title == "Part One"
        assert len(result.chapters[0].sections) == 2
        assert result.chapters[0].sections[0].title == "Section 1.1"
        assert result.chapters[0].sections[0].content == ["Content for 1.1."]
        assert result.chapters[0].sections[1].title == "Section 1.2"
        assert result.chapters[1].title == "Part Two"
        assert len(result.chapters[1].sections) == 1
        assert result.structure_depth == 2


class TestParseDocxValidationErrors:
    def test_no_headings_raises(self):
        buf = _make_docx([
            ("Normal", ""),
            ("Normal", "   "),
        ])
        with pytest.raises(ParseError, match="No headings found"):
            parse_document(buf, "docx", 1)

    def test_only_one_chapter_raises(self):
        buf = _make_docx([
            ("Heading 2", "Only Chapter"),
            ("Normal", "Some content."),
        ])
        with pytest.raises(ParseError, match="Only one chapter found"):
            parse_document(buf, "docx", 1)

    def test_depth2_no_sections_raises(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Heading 2", "Chapter Two"),
        ])
        with pytest.raises(ParseError, match="no section headings"):
            parse_document(buf, "docx", 2)

    def test_empty_chapter_title_raises(self):
        buf = _make_docx([
            ("Heading 2", "## "),
            ("Normal", "Some content."),
            ("Heading 2", "Real Chapter"),
            ("Normal", "Other content."),
        ])
        with pytest.raises(ParseError, match="empty title"):
            parse_document(buf, "docx", 1)

    def test_duplicate_chapter_titles_raises(self):
        buf = _make_docx([
            ("Heading 2", "Same Title"),
            ("Normal", "First."),
            ("Heading 2", "Same Title"),
            ("Normal", "Second."),
        ])
        with pytest.raises(ParseError, match="Duplicate chapter titles"):
            parse_document(buf, "docx", 1)

    def test_skips_empty_paragraphs(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Normal", ""),
            ("Normal", "Real content."),
            ("Normal", ""),
            ("Heading 2", "Chapter Two"),
            ("Normal", "More content."),
        ])
        result = parse_document(buf, "docx", 1)
        assert result.chapters[0].content == ["Real content."]


class TestParseDocxContentLossErrors:
    def test_body_text_before_first_chapter_raises(self):
        buf = _make_docx([
            ("Normal", "Stray intro text."),
            ("Heading 2", "Chapter One"),
            ("Normal", "First paragraph."),
            ("Heading 2", "Chapter Two"),
            ("Normal", "Second paragraph."),
        ])
        with pytest.raises(ParseError, match="before the first chapter heading"):
            parse_document(buf, "docx", 1)

    def test_body_text_between_h2_and_first_h3_raises(self):
        buf = _make_docx([
            ("Heading 2", "Part One"),
            ("Normal", "Stray text before any section."),
            ("Heading 3", "Section 1.1"),
            ("Normal", "Content for 1.1."),
            ("Heading 2", "Part Two"),
            ("Heading 3", "Section 2.1"),
            ("Normal", "Content for 2.1."),
        ])
        with pytest.raises(ParseError, match="before its first section heading"):
            parse_document(buf, "docx", 2)

    def test_h3_sections_present_when_depth_1_declared_raises(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Heading 3", "Unexpected Section"),
            ("Normal", "Some content."),
            ("Heading 2", "Chapter Two"),
            ("Normal", "More content."),
        ])
        with pytest.raises(ParseError, match="single-level structure"):
            parse_document(buf, "docx", 1)

    def test_orphan_h3_before_any_h2_raises(self):
        buf = _make_docx([
            ("Heading 3", "Orphan Section"),
            ("Normal", "Orphan content."),
            ("Heading 2", "Chapter One"),
            ("Heading 3", "Section 1.1"),
            ("Normal", "Content for 1.1."),
            ("Heading 2", "Chapter Two"),
            ("Heading 3", "Section 2.1"),
            ("Normal", "Content for 2.1."),
        ])
        with pytest.raises(ParseError, match="before any chapter heading"):
            parse_document(buf, "docx", 2)


class TestHeadingDetectionExactMatch:
    def test_custom_style_named_heading_20_is_not_a_chapter(self):
        doc = DocxDocument()
        custom_style = doc.styles.add_style("Heading 20", WD_STYLE_TYPE.PARAGRAPH)
        doc.add_paragraph("Chapter One", style="Heading 2")
        doc.add_paragraph("Looks like a heading but isn't.", style=custom_style)
        doc.add_paragraph("Real content.", style="Normal")
        doc.add_paragraph("Chapter Two", style="Heading 2")
        doc.add_paragraph("More content.", style="Normal")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = parse_document(buf, "docx", 1)
        assert len(result.chapters) == 2
        assert result.chapters[0].title == "Chapter One"
        assert "Looks like a heading but isn't." in result.chapters[0].content

    def test_markdown_prefix_in_style_structured_doc_is_not_a_chapter(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Normal", "## 3 items on the agenda"),
            ("Heading 2", "Chapter Two"),
            ("Normal", "More content."),
        ])
        result = parse_document(buf, "docx", 1)
        assert len(result.chapters) == 2
        assert result.chapters[0].title == "Chapter One"
        assert result.chapters[0].content == ["## 3 items on the agenda"]


class TestValidationGaps:
    def test_duplicate_section_titles_within_chapter_raises(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Heading 3", "Same Section"),
            ("Normal", "First."),
            ("Heading 3", "Same Section"),
            ("Normal", "Second."),
            ("Heading 2", "Chapter Two"),
            ("Heading 3", "Section A"),
            ("Normal", "Third."),
        ])
        with pytest.raises(ParseError, match="Duplicate section titles"):
            parse_document(buf, "docx", 2)

    def test_duplicate_section_titles_across_chapters_is_allowed(self):
        buf = _make_docx([
            ("Heading 2", "Chapter One"),
            ("Heading 3", "Intro"),
            ("Normal", "First."),
            ("Heading 2", "Chapter Two"),
            ("Heading 3", "Intro"),
            ("Normal", "Second."),
        ])
        result = parse_document(buf, "docx", 2)
        assert len(result.chapters) == 2

    def test_empty_title_en_raises(self):
        with pytest.raises(ParseError, match="cannot be empty"):
            build_schema(1, "   ", "ספר בדיקה")

    def test_title_en_with_slash_raises(self):
        with pytest.raises(ParseError):
            build_schema(1, "Chapter/One", "ספר בדיקה")

    def test_title_en_with_period_raises(self):
        with pytest.raises(ParseError):
            build_schema(1, "Chapter.One", "ספר בדיקה")

    def test_title_en_with_colon_raises(self):
        with pytest.raises(ParseError):
            build_schema(1, "Chapter:One", "ספר בדיקה")

    def test_title_en_with_comma_raises(self):
        with pytest.raises(ParseError):
            build_schema(1, "Chapter,One", "ספר בדיקה")

    def test_title_en_is_stripped(self):
        schema = build_schema(1, "  Test Book  ", "ספר בדיקה")
        assert schema["key"] == "Test Book"
        assert schema["titles"][0]["text"] == "Test Book"


class TestParseDocxForbiddenContent:
    def test_tables_rejected(self):
        doc = DocxDocument()
        doc.add_paragraph("Heading 2", style="Heading 2")
        doc.add_table(rows=2, cols=2)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        with pytest.raises(ParseError, match="Tables are not allowed"):
            parse_document(buf, "docx", 1)

    def test_footnotes_rejected(self):
        from docx.oxml.ns import qn
        from lxml import etree
        doc = DocxDocument()
        p = doc.add_paragraph("Some text with a footnote.")
        run = p.add_run()
        fn = etree.SubElement(run._element, qn('w:footnoteReference'))
        fn.set(qn('w:id'), '1')
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        with pytest.raises(ParseError, match="Footnotes are not allowed"):
            parse_document(buf, "docx", 1)

    def test_images_rejected(self):
        # Minimal valid 1x1 red PNG
        PNG_1x1 = bytes([
            0x89, 0x50, 0x4e, 0x47, 0x0d, 0x0a, 0x1a, 0x0a,
            0x00, 0x00, 0x00, 0x0d, 0x49, 0x48, 0x44, 0x52,
            0x00, 0x00, 0x00, 0x01, 0x00, 0x00, 0x00, 0x01,
            0x08, 0x02, 0x00, 0x00, 0x00, 0x90, 0x77, 0x53,
            0xde, 0x00, 0x00, 0x00, 0x0c, 0x49, 0x44, 0x41,
            0x54, 0x08, 0xd7, 0x63, 0xf8, 0xcf, 0xc0, 0x00,
            0x00, 0x00, 0x02, 0x00, 0x01, 0xe2, 0x21, 0xbc,
            0x33, 0x00, 0x00, 0x00, 0x00, 0x49, 0x45, 0x4e,
            0x44, 0xae, 0x42, 0x60, 0x82,
        ])
        doc = DocxDocument()
        doc.add_picture(BytesIO(PNG_1x1))
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        with pytest.raises(ParseError, match="Images are not allowed"):
            parse_document(buf, "docx", 1)
