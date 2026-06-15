from dataclasses import dataclass, field
from typing import List


class ParseError(Exception):
    pass


@dataclass
class Section:
    title: str
    content: List[str]
    word_count: int


@dataclass
class Chapter:
    title: str
    sections: List[Section]
    content: List[str]
    word_count: int


@dataclass
class ParseResult:
    chapters: List[Chapter]
    total_words: int
    structure_depth: int
    warnings: List[str] = field(default_factory=list)


def parse_document(file_obj, file_type: str, declared_depth: int) -> ParseResult:
    if file_type == "docx":
        return _parse_docx(file_obj, declared_depth)
    elif file_type == "pdf":
        raise ParseError("PDF upload is not yet supported. Please upload a .docx file.")
    else:
        raise ParseError(f"Unsupported file type: {file_type}")


def build_schema(depth: int, title_en: str, title_he: str) -> dict:
    if depth == 1:
        address_types = ["Integer"]
        section_names = ["Chapter"]
    else:
        address_types = ["Integer", "Integer"]
        section_names = ["Chapter", "Section"]

    return {
        "nodeType": "JaggedArrayNode",
        "depth": depth,
        "addressTypes": address_types,
        "sectionNames": section_names,
        "key": title_en,
        "titles": [
            {"text": title_en, "lang": "en", "primary": True},
            {"text": title_he, "lang": "he", "primary": True},
        ],
    }


def build_jagged_array(parse_result: ParseResult) -> list:
    if parse_result.structure_depth == 1:
        return ['\n'.join(ch.content) for ch in parse_result.chapters]
    else:
        return [[sec.content for sec in ch.sections] for ch in parse_result.chapters]


def _detect_heading_level(para) -> int:
    """Return 2 for H2, 3 for H3, 0 otherwise.
    First checks Word style name, then falls back to Markdown-style prefixes."""
    style_name = (para.style.name or "").lower()
    if style_name.startswith("heading 2"):
        return 2
    if style_name.startswith("heading 3"):
        return 3
    # Markdown fallback — check ### before ## to avoid false positives
    text = para.text
    if text.startswith("### "):
        return 3
    if text.startswith("## "):
        return 2
    return 0


def _check_forbidden_content(doc) -> None:
    """Raise ParseError if tables, images, or footnotes are present."""
    if doc.tables:
        raise ParseError("Tables are not allowed in uploaded documents.")

    for rel in doc.part.rels.values():
        if "image" in rel.reltype.lower():
            raise ParseError("Images are not allowed in uploaded documents.")

    try:
        from docx.oxml.ns import qn
        refs = doc.element.body.findall(f".//{qn('w:footnoteReference')}")
        if refs:
            raise ParseError("Footnotes are not allowed in uploaded documents.")
    except ImportError:
        pass


def _make_section(title: str, paragraphs: List[str]) -> Section:
    word_count = sum(len(p.split()) for p in paragraphs)
    return Section(title=title, content=paragraphs, word_count=word_count)


def _make_chapter(title: str, paragraphs: List[str], sections: List[Section], depth: int) -> Chapter:
    if depth == 2:
        word_count = sum(s.word_count for s in sections)
        return Chapter(title=title, sections=sections, content=[], word_count=word_count)
    else:
        word_count = sum(len(p.split()) for p in paragraphs)
        return Chapter(title=title, sections=[], content=paragraphs, word_count=word_count)


def _validate_chapters(chapters: List[Chapter], declared_depth: int, has_h3: bool) -> None:
    if len(chapters) == 0:
        raise ParseError("No headings found. The document must contain at least 2 chapter headings (Heading 2 style or ## prefix).")
    if len(chapters) < 2:
        raise ParseError("Only one chapter found. The document must contain at least 2 chapters.")
    for ch in chapters:
        if not ch.title.strip():
            raise ParseError("A chapter heading has an empty title.")
    titles = [ch.title for ch in chapters]
    if len(titles) != len(set(titles)):
        raise ParseError("Duplicate chapter titles found. All chapter titles must be unique.")
    if declared_depth == 2 and not has_h3:
        raise ParseError("Depth-2 structure declared but no section headings (Heading 3 / ### prefix) found.")


def _parse_docx(file_obj, declared_depth: int) -> ParseResult:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_obj)
    _check_forbidden_content(doc)

    chapters: List[Chapter] = []
    has_h3 = False

    # Current chapter state
    current_chapter_title: str = None
    current_chapter_paragraphs: List[str] = []
    current_chapter_sections: List[Section] = []

    # Current section state (for depth-2)
    current_section_title: str = None
    current_section_paragraphs: List[str] = []

    def flush_section():
        nonlocal current_section_title, current_section_paragraphs
        if current_section_title is not None:
            current_chapter_sections.append(
                _make_section(current_section_title, current_section_paragraphs)
            )
        current_section_title = None
        current_section_paragraphs = []

    def flush_chapter():
        nonlocal current_chapter_title, current_chapter_paragraphs, current_chapter_sections
        if current_chapter_title is not None:
            flush_section()
            chapters.append(
                _make_chapter(
                    current_chapter_title,
                    current_chapter_paragraphs,
                    current_chapter_sections,
                    declared_depth,
                )
            )
        current_chapter_title = None
        current_chapter_paragraphs = []
        current_chapter_sections = []

    for para in doc.paragraphs:
        level = _detect_heading_level(para)
        text = para.text

        if level == 2:
            # Strip markdown prefix if needed
            title = text[3:] if text.startswith("## ") else text
            flush_chapter()
            current_chapter_title = title

        elif level == 3:
            has_h3 = True
            title = text[4:] if text.startswith("### ") else text
            flush_section()
            current_section_title = title

        else:
            # Regular paragraph — skip if empty
            if not text.strip():
                continue
            if current_section_title is not None:
                current_section_paragraphs.append(text)
            else:
                current_chapter_paragraphs.append(text)

    flush_chapter()

    _validate_chapters(chapters, declared_depth, has_h3)

    total_words = sum(ch.word_count for ch in chapters)
    return ParseResult(
        chapters=chapters,
        total_words=total_words,
        structure_depth=declared_depth,
    )
